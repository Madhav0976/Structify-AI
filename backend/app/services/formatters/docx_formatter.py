"""
DOCX Formatter

Converts the Structured AST into a Microsoft Word Document.
Returns the binary bytes of the document.
"""
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.schemas.structure import StructuredDocument, BlockType
from .base import BaseFormatter

class DocxFormatter(BaseFormatter):
    def format(self, document: StructuredDocument) -> bytes:
        doc = Document()
        
        # Add a title if the document has a specific type
        doc.add_heading(f'Structify AI - {document.document_type.title()}', 0)
        
        for block in document.blocks:
            if block.type == BlockType.TITLE:
                p = doc.add_paragraph(block.text)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Make it look like a Title
                run = p.runs[0]
                run.font.size = Pt(24)
                run.bold = True
                
            elif block.type == BlockType.HEADING:
                # python-docx has built-in heading styles: 'Heading 1', 'Heading 2'
                doc.add_heading(block.text, level=1)
                
            elif block.type == BlockType.SUBHEADING:
                doc.add_heading(block.text, level=2)
                
            elif block.type == BlockType.LIST_ITEM:
                text = block.text
                # Strip existing markdown bullets if present so Word can use native bullets
                for c in ["- ", "* ", "• ", "1. ", "2. "]:
                    if text.startswith(c):
                        text = text[len(c):]
                        break
                doc.add_paragraph(text, style='List Bullet')
                
            elif block.type == BlockType.PARAGRAPH:
                doc.add_paragraph(block.text)
                
            else:
                doc.add_paragraph(block.text)
                
        # Save to a bytes buffer instead of a file on disk
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
